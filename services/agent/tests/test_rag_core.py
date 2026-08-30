"""步骤05.1/05.2/05.6 测试：安全抓取、格式适配器、DocumentTree、分片器。"""

from __future__ import annotations

import io

import pytest

from agent.rag.chunker import chunk_document
from agent.rag.document_tree import (
    parse_by_mime,
    parse_docx,
    parse_html,
    parse_json,
    parse_markdown_or_text,
    parse_xlsx,
)
from agent.rag.fetcher import FetchRejected, assert_host_public, assert_url_allowed, fetch, load_whitelist

WHITELIST = {"www.gov.cn", "hrss.sh.gov.cn"}


def test_whitelist_parsing():
    md = "- **www.gov.cn**\n- **hrss.sh.gov.cn**\n其他文字行"
    assert load_whitelist(md) == {"www.gov.cn", "hrss.sh.gov.cn"}


def test_url_whitelist_and_scheme():
    assert_url_allowed("https://www.gov.cn/zhengce/x.html", WHITELIST)
    with pytest.raises(FetchRejected, match="domain-not-in-whitelist"):
        assert_url_allowed("https://evil.example.com/x", WHITELIST)
    with pytest.raises(FetchRejected, match="scheme-not-allowed"):
        assert_url_allowed("file:///etc/passwd", WHITELIST)


def test_private_address_rejected(monkeypatch):
    import socket

    def fake_getaddrinfo(host, port, *args, **kwargs):
        return [(None, None, None, "", ("127.0.0.1", 0))]

    monkeypatch.setattr(socket, "getaddrinfo", fake_getaddrinfo)
    with pytest.raises(FetchRejected, match="private-address"):
        assert_host_public("www.gov.cn")


def test_fetch_rejects_oversized_content(monkeypatch):
    class FakeResponse:
        def __enter__(self):
            return self

        def __exit__(self, *a):
            return False

        status_code = 200

        def raise_for_status(self):
            return None

        url = "https://www.gov.cn/x"
        headers = {"content-type": "text/html"}

        def iter_bytes(self):
            yield b"x" * (6 * 1024 * 1024)

    class FakeClient:
        def __init__(self, *a, **k):
            pass

        def __enter__(self):
            return self

        def __exit__(self, *a):
            return False

        def stream(self, method, url):
            return FakeResponse()

    monkeypatch.setattr("agent.rag.fetcher.httpx.Client", FakeClient)
    with pytest.raises(FetchRejected, match="payload-too-large"):
        fetch("https://www.gov.cn/x", WHITELIST, max_bytes=5 * 1024 * 1024)


def test_redirect_outside_whitelist_rejected(monkeypatch):
    class FakeResponse:
        def __enter__(self):
            return self

        def __exit__(self, *a):
            return False

        status_code = 302
        url = "https://www.gov.cn/x"
        headers = {"location": "https://evil.example.com/final"}

        def iter_bytes(self):
            yield b""

    class FakeClient:
        def __init__(self, *a, **k):
            pass

        def __enter__(self):
            return self

        def __exit__(self, *a):
            return False

        def stream(self, method, url):
            return FakeResponse()

    monkeypatch.setattr("agent.rag.fetcher.httpx.Client", FakeClient)
    with pytest.raises(FetchRejected, match="domain-not-in-whitelist"):
        fetch("https://www.gov.cn/x", WHITELIST, max_bytes=1024)


def test_format_adapters_markdown_html_json():
    md = ("# 第一章 总则\n第一条 缴费年限满15年。\n相关说明文字。").encode()
    md_result = parse_markdown_or_text(md)
    assert md_result.tree.children[0].type == "chapter"
    assert md_result.pipeline_version == "rag-parse-v1"

    html = ("<html><body><h1>上海市社保政策</h1><p>条款内容。</p></body></html>").encode()
    html_result = parse_html(html)
    assert html_result.tree.children[0].type == "chapter"

    data = b'{"rule_id": "R-010", "priority": 10}'
    json_result = parse_json(data)
    assert any("R-010" in c.text for c in json_result.tree.children)


def test_format_adapters_docx_xlsx():
    from docx import Document
    from openpyxl import Workbook

    doc = Document()
    doc.add_heading("上海政策文件", level=1)
    doc.add_paragraph("正文段落。")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "项目"
    table.rows[1].cells[0].text = "数值"
    buf = io.BytesIO()
    doc.save(buf)
    docx_result = parse_docx(buf.getvalue())
    assert docx_result.tree.children[0].type == "chapter"
    assert any(c.type == "table" for c in docx_result.tree.children)

    wb = Workbook()
    ws = wb.active
    ws.title = "参数表"
    ws.append(["参数", "值"])
    ws.append(["最低工资", 2690])
    buf2 = io.BytesIO()
    wb.save(buf2)
    xlsx_result = parse_xlsx(buf2.getvalue())
    assert xlsx_result.tree.children[0].type == "chapter"
    assert any("2690" in c.text for c in xlsx_result.tree.children[0].children)

    # 按扩展名路由
    routed = parse_by_mime("application/vnd.openxmlformats-officedocument.wordprocessingml.document", "a.docx", buf.getvalue())
    assert routed.tree.type == "document"


def test_chunker_parent_child_and_table():
    from agent.rag.document_tree import TreeNode

    md = ("# 第一章\n第一条 内容A。\n# 第二章\n表格见附件。").encode()
    tree = parse_markdown_or_text(md).tree
    chunks = chunk_document(tree, "doc-v1")
    assert chunks, "至少产生一个 chunk"
    ids = {c.chunk_id for c in chunks}
    for c in chunks:
        if c.parent_chunk_id:
            assert c.parent_chunk_id in ids

    # 表格分片：每组复制表头。
    table_node = TreeNode(type="table", children=[
        TreeNode(type="row", text="项目 | 数值"),
        TreeNode(type="row", text="最低工资 | 2690"),
        TreeNode(type="row", text="封顶线 | 36549"),
    ])
    root = TreeNode(type="document", children=[TreeNode(type="chapter", text="附表", children=[table_node])])
    table_chunks = chunk_document(root, "doc-v2")
    assert table_chunks
