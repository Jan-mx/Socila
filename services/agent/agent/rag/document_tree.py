"""DocumentTree 模型与格式适配器（RAG-FR-004～007，步骤05.2/05.4）。

- 统一映射到 DocumentTree（节点含类型、文本、页码、子节点、表格）。
- 适配器：HTML(httpx+lxml)、DOCX(python-docx)、XLSX(openpyxl 只读)、JSON、Markdown/TXT（行式）。
- 解析器版本随 DocumentVersion 记录（pipelineVersion）。
"""

from __future__ import annotations

import json
from dataclasses import dataclass, field
from typing import Any


@dataclass
class TreeNode:
    type: str  # document|chapter|section|article|paragraph|list|table|row
    text: str = ""
    page: int | None = None
    children: list["TreeNode"] = field(default_factory=list)
    meta: dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> dict[str, Any]:
        return {
            "type": self.type,
            "text": self.text,
            "page": self.page,
            "meta": self.meta,
            "children": [c.to_dict() for c in self.children],
        }


@dataclass
class ParseResult:
    tree: TreeNode
    pipeline_version: str
    warnings: list[str] = field(default_factory=list)


PIPELINE_VERSION = "rag-parse-v1"


def parse_markdown_or_text(raw: bytes) -> ParseResult:
    text = raw.decode("utf-8", errors="replace")
    root = TreeNode(type="document")
    current_chapter: TreeNode | None = None
    current_article: TreeNode | None = None
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("# "):
            current_chapter = TreeNode(type="chapter", text=stripped[2:])
            root.children.append(current_chapter)
            current_article = None
        elif stripped.startswith("第") and ("条" in stripped[:8] or "章" in stripped[:8]):
            current_article = TreeNode(type="article", text=stripped)
            (current_chapter or root).children.append(current_article)
        else:
            node = TreeNode(type="paragraph", text=stripped)
            (current_article or current_chapter or root).children.append(node)
    return ParseResult(tree=root, pipeline_version=PIPELINE_VERSION)


def parse_json(raw: bytes) -> ParseResult:
    data = json.loads(raw.decode("utf-8", errors="replace"))
    root = TreeNode(type="document")

    def walk(obj: Any, parent: TreeNode) -> None:
        if isinstance(obj, dict):
            for k, v in obj.items():
                node = TreeNode(type="paragraph", text=f"{k}: ")
                if isinstance(v, (dict, list)):
                    parent.children.append(node)
                    walk(v, node)
                else:
                    node.text = f"{k}: {v}"
                    parent.children.append(node)
        elif isinstance(obj, list):
            for item in obj:
                if isinstance(item, (dict, list)):
                    walk(item, parent)
                else:
                    parent.children.append(TreeNode(type="paragraph", text=str(item)))

    walk(data, root)
    return ParseResult(tree=root, pipeline_version=PIPELINE_VERSION)


def parse_html(raw: bytes) -> ParseResult:
    from lxml import html as lxml_html

    doc = lxml_html.fromstring(raw.decode("utf-8", errors="replace"))
    root = TreeNode(type="document")
    for heading in doc.xpath("//h1|//h2|//h3"):
        chapter = TreeNode(type="chapter", text=heading.text_content().strip())
        root.children.append(chapter)
        for sibling in heading.itersiblings():
            tag = sibling.tag.lower()
            if tag in ("h1", "h2", "h3"):
                break
            if tag in ("p", "li"):
                text = sibling.text_content().strip()
                if text:
                    chapter.children.append(TreeNode(type="paragraph", text=text))
    if not root.children:
        text = " ".join(doc.text_content().split())
        root.children.append(TreeNode(type="paragraph", text=text))
    return ParseResult(tree=root, pipeline_version=PIPELINE_VERSION)


def parse_docx(raw: bytes) -> ParseResult:
    import io

    from docx import Document

    document = Document(io.BytesIO(raw))
    root = TreeNode(type="document")
    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if para.style.name.lower().startswith("heading"):
            root.children.append(TreeNode(type="chapter", text=text))
        else:
            root.children.append(TreeNode(type="paragraph", text=text))
    for table in document.tables:
        tnode = TreeNode(type="table")
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            tnode.children.append(TreeNode(type="row", text=" | ".join(cells)))
        root.children.append(tnode)
    return ParseResult(tree=root, pipeline_version=PIPELINE_VERSION)


def parse_xlsx(raw: bytes) -> ParseResult:
    import io

    from openpyxl import load_workbook

    workbook = load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
    root = TreeNode(type="document")
    for sheet in workbook.worksheets:
        sheet_node = TreeNode(type="chapter", text=sheet.title)
        for row in sheet.iter_rows(values_only=True):
            cells = ["" if v is None else str(v) for v in row]
            if any(cells):
                sheet_node.children.append(TreeNode(type="row", text=" | ".join(cells)))
        root.children.append(sheet_node)
    return ParseResult(tree=root, pipeline_version=PIPELINE_VERSION)


def parse_by_mime(mime: str, filename: str, raw: bytes) -> ParseResult:
    """按 MIME/扩展名路由（RAG-FR-001）。不支持/超限由调用方拒绝。"""
    lower = filename.lower()
    if "pdf" in mime or lower.endswith(".pdf"):
        raise NotImplementedError("pdf routed to 05.3 OCR pipeline")
    if "wordprocessingml" in mime or lower.endswith(".docx"):
        return parse_docx(raw)
    if "spreadsheetml" in mime or lower.endswith(".xlsx"):
        return parse_xlsx(raw)
    if "html" in mime or lower.endswith((".html", ".htm")):
        return parse_html(raw)
    if "json" in mime or lower.endswith(".json"):
        return parse_json(raw)
    if lower.endswith((".md", ".markdown", ".txt")) or mime.startswith("text/"):
        return parse_markdown_or_text(raw)
    raise NotImplementedError(f"unsupported mime: {mime}")


# ── 资源限制与流式（G4：XLSX 10万行、JSON>5MB 流式）──────────────────────────


def parse_xlsx_with_limits(raw: bytes, max_rows: int = 100_000) -> ParseResult:
    from io import BytesIO

    from openpyxl import load_workbook

    workbook = load_workbook(BytesIO(raw), read_only=True, data_only=True)
    total_rows = 0
    root = TreeNode(type="document")
    for sheet in workbook.worksheets:
        sheet_node = TreeNode(type="chapter", text=sheet.title)
        for row in sheet.iter_rows(values_only=True):
            total_rows += 1
            if total_rows > max_rows:
                raise ValueError("row-limit-exceeded")
            cells = ["" if v is None else str(v) for v in row]
            if any(cells):
                sheet_node.children.append(TreeNode(type="row", text=" | ".join(cells)))
        root.children.append(sheet_node)
    return ParseResult(tree=root, pipeline_version=PIPELINE_VERSION)


def parse_json_streamed(raw: bytes) -> ParseResult:
    """≤5MB 直接解析；>5MB 用 ijson 流式（PRD §9.2 / operational-baseline）。"""
    if len(raw) <= 5 * 1024 * 1024:
        return parse_json(raw)
    import ijson

    root = TreeNode(type="document")
    count = 0
    for item in ijson.items(io := __import__("io").BytesIO(raw), "items.item"):
        count += 1
        if isinstance(item, dict):
            node = TreeNode(type="paragraph", text=json.dumps(item, ensure_ascii=False, sort_keys=True))
            root.children.append(node)
        else:
            root.children.append(TreeNode(type="paragraph", text=str(item)))
    if count == 0:
        return parse_json(raw)
    return ParseResult(tree=root, pipeline_version=PIPELINE_VERSION + "+streamed")
